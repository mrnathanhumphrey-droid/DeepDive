import os
import re
import io
import json
import glob as globlib
from dataclasses import asdict
from datetime import datetime
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from agents.orchestrator import Orchestrator
from models.schemas import AgentType
from vector_store import CorrectionStore
from google_drive import upload_docx_to_drive, is_drive_configured, is_drive_authenticated

RESULTS_DIR = os.path.join(os.path.dirname(__file__), "results")
LOG_DIR = os.path.join(os.path.dirname(__file__), "results")

# Page config
st.set_page_config(
    page_title="DeepDive - Multi-Agent Research & Analysis",
    page_icon="🔍",
    layout="wide",
)

# ── Neon Punk CSS ──────────────────────────────────────────────────
st.markdown("""
<style>
    /* ── Import fonts ── */
    @import url('https://fonts.googleapis.com/css2?family=Permanent+Marker&family=Space+Mono:wght@400;700&display=swap');

    /* ── CSS Custom Properties ── */
    :root {
        --neon-green:   #39ff14;
        --neon-cyan:    #00f0ff;
        --neon-magenta: #ff2d95;
        --neon-orange:  #ff6b2b;
        --neon-purple:  #8b00ff;
        --neon-yellow:  #ffff00;
        --black:        #0a0a12;
        --black-mid:    #0f0f1a;
        --black-soft:   #161624;
        --gray-dark:    #1e1e30;
        --gray-mid:     #2a2a3d;
        --gray-muted:   #666666;
        --gray-light:   #999999;
        --white:        #f0f0f0;
        --glow-green:   rgba(57,255,20,.15);
        --glow-cyan:    rgba(0,240,255,.15);
        --glow-magenta: rgba(255,45,149,.15);
        --shadow-neon:  0 0 20px rgba(57,255,20,.3), 0 0 60px rgba(57,255,20,.1);
        --shadow-magenta: 0 0 20px rgba(255,45,149,.3), 0 0 60px rgba(255,45,149,.1);
        --shadow-cyan:  0 0 20px rgba(0,240,255,.3), 0 0 60px rgba(0,240,255,.1);
        --transition:   0.22s cubic-bezier(.4,0,.2,1);
    }

    /* ── Global overrides ── */
    .stApp {
        background: var(--black) !important;
        font-family: 'Space Mono', monospace !important;
    }

    /* Film grain overlay */
    .stApp::after {
        content: '';
        position: fixed;
        inset: 0;
        background-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 256 256' xmlns='http://www.w3.org/2000/svg'%3E%3Cfilter id='n'%3E%3CfeTurbulence type='fractalNoise' baseFrequency='.85' numOctaves='4' stitchTiles='stitch'/%3E%3C/filter%3E%3Crect width='100%25' height='100%25' filter='url(%23n)' opacity='.04'/%3E%3C/svg%3E");
        pointer-events: none;
        z-index: 9999;
        opacity: 0.4;
    }

    /* Scanline overlay */
    .stApp::before {
        content: '';
        position: fixed;
        inset: 0;
        background: repeating-linear-gradient(
            0deg,
            transparent 0px,
            transparent 2px,
            rgba(0,0,0,.06) 2px,
            rgba(0,0,0,.06) 4px
        );
        pointer-events: none;
        z-index: 9998;
    }

    /* ── Typography ── */
    h1, h2, h3 {
        font-family: 'Permanent Marker', cursive !important;
        line-height: 1.1 !important;
    }

    h1 {
        color: var(--neon-magenta) !important;
        text-shadow: 0 0 30px rgba(255,45,149,.5), 0 0 80px rgba(255,45,149,.2);
        font-size: clamp(2rem, 5vw, 3.5rem) !important;
    }

    h2 {
        color: var(--neon-green) !important;
        text-shadow: 0 0 15px rgba(57,255,20,.3);
        font-size: clamp(1.3rem, 3vw, 2rem) !important;
    }

    h3 {
        color: var(--neon-cyan) !important;
        text-shadow: 0 0 10px rgba(0,240,255,.3);
        font-size: clamp(1rem, 2vw, 1.4rem) !important;
    }

    p, li, span, div {
        font-family: 'Space Mono', monospace !important;
    }

    /* ── Scrollbar ── */
    ::-webkit-scrollbar { width: 8px; }
    ::-webkit-scrollbar-track { background: var(--black); }
    ::-webkit-scrollbar-thumb { background: var(--gray-mid); border-radius: 4px; }
    ::-webkit-scrollbar-thumb:hover { background: var(--neon-green); }

    /* ── Sidebar ── */
    [data-testid="stSidebar"] {
        background: var(--black-mid) !important;
        border-right: 1px solid var(--gray-dark) !important;
    }

    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2 {
        color: var(--neon-magenta) !important;
    }

    /* ── Agent cards ── */
    .agent-card {
        background: var(--black-soft);
        border: 1px solid var(--gray-dark);
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
        transition: transform var(--transition), box-shadow var(--transition), border-color var(--transition);
    }

    .agent-card:hover {
        transform: translateY(-3px);
        border-color: var(--neon-green);
        box-shadow: var(--shadow-neon);
    }

    .confidence-high {
        border-left: 4px solid var(--neon-green) !important;
        box-shadow: inset 4px 0 12px rgba(57,255,20,.1);
    }
    .confidence-med {
        border-left: 4px solid var(--neon-orange) !important;
        box-shadow: inset 4px 0 12px rgba(255,107,43,.1);
    }
    .confidence-low {
        border-left: 4px solid var(--neon-magenta) !important;
        box-shadow: inset 4px 0 12px rgba(255,45,149,.1);
    }

    /* ── Progress bar ── */
    .stProgress > div > div > div {
        height: 8px;
        background: linear-gradient(90deg, var(--neon-green), var(--neon-cyan)) !important;
        box-shadow: 0 0 12px rgba(57,255,20,.4);
    }

    /* ── Tangential box ── */
    .tangential-box {
        background: var(--black-soft);
        border: 1px solid var(--gray-dark);
        border-radius: 8px;
        padding: 12px;
        margin-top: 8px;
        transition: border-color var(--transition);
    }

    .tangential-box:hover {
        border-color: var(--neon-purple);
        box-shadow: 0 0 15px rgba(139,0,255,.2);
    }

    /* ── Tabs ── */
    .stTabs [data-baseweb="tab-list"] {
        gap: 4px;
        background: var(--black-mid);
        border-radius: 8px;
        padding: 4px;
    }

    .stTabs [data-baseweb="tab"] {
        font-family: 'Space Mono', monospace !important;
        font-size: 0.78rem !important;
        letter-spacing: 0.05em;
        color: var(--gray-light) !important;
        background: transparent !important;
        border: 1px solid transparent !important;
        border-radius: 6px !important;
        padding: 8px 16px !important;
        transition: all var(--transition) !important;
    }

    .stTabs [data-baseweb="tab"]:hover {
        color: var(--neon-green) !important;
        text-shadow: 0 0 8px rgba(57,255,20,.5);
        background: rgba(57,255,20,.05) !important;
    }

    .stTabs [aria-selected="true"] {
        color: var(--neon-green) !important;
        background: rgba(57,255,20,.1) !important;
        border-color: var(--neon-green) !important;
        box-shadow: 0 0 15px rgba(57,255,20,.2) !important;
    }

    /* ── Buttons ── */
    .stButton > button {
        font-family: 'Space Mono', monospace !important;
        font-size: 0.85rem !important;
        letter-spacing: 0.08em;
        text-transform: uppercase;
        background: var(--neon-green) !important;
        color: var(--black) !important;
        border: 1.5px solid var(--neon-green) !important;
        border-radius: 4px !important;
        box-shadow: 0 0 15px rgba(57,255,20,.3);
        transition: all var(--transition) !important;
        font-weight: 700 !important;
    }

    .stButton > button:hover {
        background: transparent !important;
        color: var(--neon-green) !important;
        box-shadow: 0 0 30px rgba(57,255,20,.5) !important;
        transform: translateY(-2px);
    }

    .stButton > button:active {
        transform: translateY(0);
    }

    /* ── Text input ── */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        background: var(--black-mid) !important;
        border: 1px solid var(--gray-mid) !important;
        border-radius: 4px !important;
        color: var(--white) !important;
        font-family: 'Space Mono', monospace !important;
        font-size: 0.85rem !important;
        transition: border-color var(--transition), box-shadow var(--transition) !important;
    }

    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: var(--neon-cyan) !important;
        box-shadow: 0 0 12px rgba(0,240,255,.15) !important;
    }

    /* ── Expanders ── */
    .streamlit-expanderHeader {
        font-family: 'Space Mono', monospace !important;
        font-size: 0.85rem !important;
        color: var(--gray-light) !important;
        background: var(--black-soft) !important;
        border: 1px solid var(--gray-dark) !important;
        border-radius: 6px !important;
        transition: all var(--transition) !important;
    }

    .streamlit-expanderHeader:hover {
        color: var(--neon-cyan) !important;
        border-color: var(--neon-cyan) !important;
        text-shadow: 0 0 8px rgba(0,240,255,.4);
    }

    /* ── Dividers ── */
    hr {
        border-color: var(--gray-dark) !important;
        box-shadow: 0 0 8px rgba(57,255,20,.05);
    }

    /* ── Info/Warning/Error boxes ── */
    .stAlert [data-testid="stMarkdownContainer"] {
        font-family: 'Space Mono', monospace !important;
    }

    div[data-testid="stNotification"] {
        background: var(--black-soft) !important;
        border: 1px solid var(--gray-dark) !important;
        border-radius: 6px !important;
    }

    /* ── Metrics ── */
    [data-testid="stMetricValue"] {
        font-family: 'Permanent Marker', cursive !important;
        color: var(--neon-green) !important;
        text-shadow: 0 0 10px rgba(57,255,20,.4);
    }

    [data-testid="stMetricDelta"] {
        font-family: 'Space Mono', monospace !important;
    }

    /* ── Download button ── */
    .stDownloadButton > button {
        font-family: 'Space Mono', monospace !important;
        background: transparent !important;
        color: var(--neon-cyan) !important;
        border: 1.5px solid var(--neon-cyan) !important;
        box-shadow: none !important;
        letter-spacing: 0.08em;
        text-transform: uppercase;
    }

    .stDownloadButton > button:hover {
        background: rgba(0,240,255,.1) !important;
        box-shadow: 0 0 20px rgba(0,240,255,.3) !important;
        transform: translateY(-2px);
    }

    /* ── Selectbox / Dropdown ── */
    .stSelectbox > div > div {
        background: var(--black-mid) !important;
        border-color: var(--gray-mid) !important;
        font-family: 'Space Mono', monospace !important;
    }

    /* ── Citation status badges ── */
    .badge-verified {
        display: inline-block;
        background: rgba(57,255,20,.15);
        color: var(--neon-green);
        padding: 2px 8px;
        border-radius: 4px;
        font-size: 0.75rem;
        font-family: 'Space Mono', monospace;
        letter-spacing: 0.05em;
    }
    .badge-unverified {
        display: inline-block;
        background: rgba(0,240,255,.15);
        color: var(--neon-cyan);
        padding: 2px 8px;
        border-radius: 4px;
        font-size: 0.75rem;
        font-family: 'Space Mono', monospace;
        letter-spacing: 0.05em;
    }
    .badge-fabrication {
        display: inline-block;
        background: rgba(255,45,149,.15);
        color: var(--neon-magenta);
        padding: 2px 8px;
        border-radius: 4px;
        font-size: 0.75rem;
        font-family: 'Space Mono', monospace;
        letter-spacing: 0.05em;
    }

    /* ── Glitch animation ── */
    @keyframes glitch {
        0%   { text-shadow: 2px 0 #00f0ff, -2px 0 #ff2d95; }
        25%  { text-shadow: -2px -1px #ff2d95, 2px 1px #39ff14; }
        50%  { text-shadow: 1px 2px #39ff14, -1px -2px #00f0ff; }
        75%  { text-shadow: -1px 1px #00f0ff, 1px -1px #ff6b2b; }
        100% { text-shadow: 2px 0 #00f0ff, -2px 0 #ff2d95; }
    }

    @keyframes pulse-glow {
        0%, 100% { opacity: 0.6; }
        50% { opacity: 1; }
    }

    @keyframes fade-up {
        from { opacity: 0; transform: translateY(20px); }
        to   { opacity: 1; transform: translateY(0); }
    }

    /* ── Title animation ── */
    .deepdive-title {
        font-family: 'Permanent Marker', cursive;
        font-size: clamp(2.5rem, 6vw, 4rem);
        color: var(--neon-magenta);
        text-shadow: 0 0 30px rgba(255,45,149,.5), 0 0 80px rgba(255,45,149,.2);
        animation: fade-up 0.8s ease both;
        margin-bottom: 0;
        line-height: 1.1;
    }

    .deepdive-title:hover {
        animation: glitch 0.3s linear infinite;
    }

    .deepdive-subtitle {
        font-family: 'Space Mono', monospace;
        font-size: 0.78rem;
        letter-spacing: 0.15em;
        color: var(--neon-cyan);
        text-shadow: 0 0 10px rgba(0,240,255,.4);
        text-transform: uppercase;
        animation: fade-up 0.8s 0.1s ease both;
    }

    /* ── Agent progress cards ── */
    .agent-progress {
        background: var(--black-soft);
        border: 1px solid var(--gray-dark);
        border-radius: 6px;
        padding: 12px 16px;
        margin: 6px 0;
        display: flex;
        align-items: center;
        gap: 12px;
        transition: all var(--transition);
    }

    .agent-progress:hover {
        border-color: var(--neon-green);
        box-shadow: 0 0 12px rgba(57,255,20,.15);
    }

    .agent-progress .agent-icon {
        font-size: 1.4rem;
    }

    .agent-progress .agent-name {
        font-family: 'Space Mono', monospace;
        font-size: 0.8rem;
        letter-spacing: 0.05em;
        color: var(--white);
        flex: 1;
    }

    .agent-progress .agent-conf {
        font-family: 'Space Mono', monospace;
        font-size: 0.75rem;
        padding: 2px 10px;
        border-radius: 4px;
    }

    .agent-progress .conf-high {
        background: rgba(57,255,20,.15);
        color: var(--neon-green);
    }
    .agent-progress .conf-med {
        background: rgba(255,107,43,.15);
        color: var(--neon-orange);
    }
    .agent-progress .conf-low {
        background: rgba(255,45,149,.15);
        color: var(--neon-magenta);
    }

    /* ── Section headers with neon accent ── */
    .section-header {
        border-bottom: 1px solid var(--gray-dark);
        padding-bottom: 8px;
        margin-bottom: 16px;
    }

    .section-header h2 {
        display: inline-block;
        position: relative;
    }

    .section-header h2::after {
        content: '';
        position: absolute;
        bottom: -9px;
        left: 0;
        width: 60px;
        height: 2px;
        background: var(--neon-green);
        box-shadow: 0 0 8px rgba(57,255,20,.5);
    }

    /* ── Focus styles ── */
    :focus-visible {
        outline: 2px solid var(--neon-green) !important;
        outline-offset: 3px;
        border-radius: 4px;
    }

    /* ── Link styling ── */
    a {
        color: var(--neon-cyan) !important;
        text-decoration: none !important;
        transition: all var(--transition);
    }

    a:hover {
        color: var(--neon-green) !important;
        text-shadow: 0 0 8px rgba(57,255,20,.5);
    }

    /* ── Markdown code blocks ── */
    code {
        background: var(--black-mid) !important;
        border: 1px solid var(--gray-dark) !important;
        color: var(--neon-orange) !important;
        font-family: 'Space Mono', monospace !important;
        padding: 2px 6px !important;
        border-radius: 3px !important;
    }

    pre {
        background: var(--black-mid) !important;
        border: 1px solid var(--gray-dark) !important;
        border-radius: 6px !important;
    }
</style>
""", unsafe_allow_html=True)

AGENT_ICONS = {
    AgentType.FACTS: "📋",
    AgentType.CONTEXT: "📚",
    AgentType.CONTEXT_US: "🇺🇸",
    AgentType.CONTEXT_WORLD: "🌍",
    AgentType.CONTEXT_SYNTHESIS: "📚",
    AgentType.PERSPECTIVES: "🔄",
    AgentType.TIMELINE: "⏱️",
    AgentType.SPLIT_REVIEWER: "✅",
    AgentType.FACT_CHECKER: "🔎",
    AgentType.RESEARCH_REVIEW: "🎓",
    AgentType.GOVERNMENT_DOCS: "🏛️",
    AgentType.META_REVIEW: "🧠",
    AgentType.PROMPT_ENGINEER: "🔬",
    AgentType.ECONOMICS_DATA: "📊",
    AgentType.ECONOMICS_POLICY: "💹",
    # Historical pipeline agents
    AgentType.HISTORICAL_ANCHOR: "🏛️",
    AgentType.ERA_CONTEXT: "🕰️",
    AgentType.PRIMARY_SOURCE: "📜",
    AgentType.CAUSAL_CHAIN: "🔗",
    AgentType.MODERN_IMPACT: "📡",
    AgentType.SCHOLARLY_CONSENSUS: "🎓",
    AgentType.COUNTERFACTUAL: "🔀",
    AgentType.RIPPLE_TIMELINE: "🌊",
}

# ── Session state initialization ────────────────────────────────────

def init_session_state():
    defaults = {
        "pipeline_state": None,
        "phase": "input",
        "correction_round": 0,
        "orchestrator": None,
        "_topic": "",
        "_user_feedback": "",
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val


# ── Rendering helpers ───────────────────────────────────────────────

def render_home_button():
    """Permanent home button at the top of every page."""
    if st.session_state.phase != "input":
        if st.button("🏠 New Analysis", key="home_btn_top"):
            st.session_state.phase = "input"
            st.session_state.pipeline_state = None
            st.session_state.correction_round = 0
            st.rerun()


def render_header():
    st.title("DeepDive")
    st.caption("Multi-Agent Research & Analysis — Break down any topic through multiple expert lenses")
    render_home_button()
    st.divider()


def render_input():
    st.markdown(
        '<div style="text-align:center;padding:2rem 0 0.5rem;">'
        '<p class="deepdive-subtitle">multi-agent research & analysis</p>'
        '<h1 class="deepdive-title">DEEPDIVE</h1>'
        '</div>',
        unsafe_allow_html=True,
    )
    col1, col2 = st.columns([4, 1])
    with col1:
        topic = st.text_input(
            "Enter a topic, headline, or historical question",
            placeholder="e.g., EU AI Act, Chevron deference reversal, legacy of the Voting Rights Act",
            key="topic_input",
        )
    with col2:
        st.write("")
        st.write("")
        analyze = st.button("Analyze", type="primary", use_container_width=True)
    return topic, analyze


def render_confidence(confidence):
    confidence_pct = confidence * 100
    if confidence_pct >= 70:
        color = "#39ff14"
        cls = "conf-high"
    elif confidence_pct >= 50:
        color = "#ff6b2b"
        cls = "conf-med"
    else:
        color = "#ff2d95"
        cls = "conf-low"
    return f"<span style='color:{color};font-family:Space Mono,monospace;font-size:0.85rem;'>{confidence_pct:.0f}%</span>"


def render_sub_topics(sub_topics, tangential_topics=None):
    st.subheader("Topic Breakdown")
    cols = st.columns(len(sub_topics))
    for i, st_topic in enumerate(sub_topics):
        with cols[i]:
            icon = AGENT_ICONS.get(st_topic.agent_type, "🔍")
            st.markdown(f"### {icon} {st_topic.agent_type.value.replace('_', ' ').title()}")
            st.markdown(f"**{st_topic.title}**")
            st.caption(st_topic.description)
            if st_topic.keywords:
                st.markdown(" ".join(f"`{kw}`" for kw in st_topic.keywords))

            # Tangential topics
            tangentials = (tangential_topics or {}).get(st_topic.agent_type.value, [])
            if tangentials:
                with st.expander("Tangential Topics", expanded=False):
                    for t in tangentials:
                        summary = t.get("summary", "")
                        search_term = t.get("search_term", "")
                        st.markdown(f"{summary}")
                        if search_term:
                            st.markdown(
                                f"🔎 **[{search_term}](https://news.google.com/search?q={search_term.replace(' ', '+')})**"
                            )
                        st.markdown("---")


def parse_citation_block(content):
    """Extract the JSON citation block from agent output, return (prose, citation_data)."""
    import json as _json

    # Look for ```json ... ``` block
    idx = content.rfind("```json")
    if idx == -1:
        # Try bare JSON at end starting with {"agent"
        idx = content.rfind('{"agent"')
        if idx == -1:
            return content, None
        prose = content[:idx].rstrip()
        raw_json = content[idx:]
        # Find closing brace
        brace_depth = 0
        end = 0
        for i, ch in enumerate(raw_json):
            if ch == '{':
                brace_depth += 1
            elif ch == '}':
                brace_depth -= 1
                if brace_depth == 0:
                    end = i + 1
                    break
        raw_json = raw_json[:end]
    else:
        prose = content[:idx].rstrip()
        end_idx = content.find("```", idx + 7)
        if end_idx == -1:
            raw_json = content[idx + 7:]
        else:
            raw_json = content[idx + 7:end_idx]

    try:
        data = _json.loads(raw_json.strip())
        return prose, data
    except (_json.JSONDecodeError, ValueError):
        return content, None


def render_citation_block(citation_data):
    """Render a parsed citation block as formatted text with status icons."""
    if not citation_data:
        return

    STATUS_ICONS = {
        "verified": "✅",
        "plausible_unverified": "🟡",
        "fabrication_risk": "🔴",
    }

    with st.expander("Citation Verification Report", expanded=False):
        # Verified citations
        citations = citation_data.get("citations", [])
        verified_citations = citation_data.get("verified_citations", [])
        all_cites = citations or verified_citations

        if all_cites:
            st.markdown("**Verified Sources:**")
            for c in all_cites:
                if isinstance(c, str):
                    st.markdown(f"✅ {c}")
                    continue
                status = c.get("status", "verified")
                icon = STATUS_ICONS.get(status, "❓")
                summary = (
                    c.get("claim_summary", "")
                    or c.get("event_summary", "")
                    or c.get("metric", "")
                )
                source = (
                    c.get("source_title", "")
                    or c.get("attributed_to", "")
                    or c.get("publication", "")
                    or c.get("publishing_agency", "")
                    or c.get("document_title", "")
                )
                date = (
                    c.get("source_date", "")
                    or c.get("event_date", "")
                    or c.get("year", "")
                    or c.get("reference_period", "")
                    or c.get("document_date", "")
                )
                st.markdown(f"{icon} **{summary}**")
                details = []
                if source:
                    details.append(f"Source: {source}")
                if date:
                    details.append(f"Date: {date}")
                if c.get("identifier") and c["identifier"] != "none":
                    confirmed = "confirmed" if c.get("identifier_confirmed") else "unconfirmed"
                    details.append(f"ID: `{c['identifier']}` ({confirmed})")
                if c.get("source_tier"):
                    details.append(f"Tier: {c['source_tier']}")
                if details:
                    st.caption(" | ".join(details))

        # Upstream hedged claims
        hedged = citation_data.get("upstream_hedged_claims", [])
        if hedged:
            st.markdown("**Hedged Claims (unverified upstream):**")
            for h in hedged:
                if isinstance(h, str):
                    st.markdown(f"🟡 {h}")
                elif isinstance(h, dict):
                    st.markdown(f"🟡 {h.get('claim_summary', '')}")
                    st.caption(h.get("reason", ""))

        # Downgraded citations
        downgraded = citation_data.get("downgraded_citations", [])
        if downgraded:
            st.markdown("**Downgraded Citations:**")
            for d in downgraded:
                if isinstance(d, str):
                    st.markdown(f"🟡 {d}")
                elif isinstance(d, dict):
                    st.markdown(
                        f"🟡 **{d.get('claim_summary', '')}** — "
                        f"{d.get('original_status', '')} -> {d.get('new_status', '')}"
                    )
                    st.caption(f"Reason: {d.get('reason', '')} | From: {d.get('upstream_agent', '')}")

        # Fabrication risk flags
        fab_flags = citation_data.get("fabrication_risk_flags", [])
        if fab_flags:
            st.markdown("**Fabrication Risk Flags:**")
            for f in fab_flags:
                if isinstance(f, str):
                    st.markdown(f"🔴 **{f}**")
                elif isinstance(f, dict):
                    st.markdown(f"🔴 **{f.get('claim_summary', '')}**")
                    st.caption(f"Reason: {f.get('reason', '')} | From: {f.get('upstream_agent', '')}")

        # Removed claims
        removed = citation_data.get("removed_claims", [])
        if removed:
            st.markdown("**Removed Claims:**")
            for r in removed:
                if isinstance(r, str):
                    st.markdown(f"🔴 ~~{r}~~")
                elif isinstance(r, dict):
                    summary = r.get('claim_summary', r.get('event_summary', r.get('metric', 'Unknown')))
                    st.markdown(f"🔴 ~~{summary}~~")
                    reason = r.get('reason', '')
                    if reason:
                        st.caption(f"Reason: {reason}")

        # Cross-agent conflicts
        conflicts = citation_data.get("cross_agent_conflicts", [])
        if conflicts:
            st.markdown("**Cross-Agent Conflicts:**")
            for c in conflicts:
                if isinstance(c, str):
                    st.markdown(f"⚠️ {c}")
                elif isinstance(c, dict):
                    st.markdown(f"⚠️ **{c.get('claim_topic', '')}**")
                    st.caption(
                        f"{c.get('conflict_description', '')} | "
                        f"Higher-tier: {c.get('higher_tier_source', '')}"
                    )

        # Summary counts
        total_verified = len([c for c in all_cites if c.get("status") == "verified"])
        total_unverified = len([c for c in all_cites if c.get("status") == "plausible_unverified"])
        total_removed = len(removed) + len(fab_flags)
        st.divider()
        st.markdown(
            f"**Summary:** {total_verified} ✅ verified | "
            f"{total_unverified} 🟡 unverified | "
            f"{total_removed} 🔴 removed/flagged"
        )


def render_agent_results(results):
    st.subheader("Agent Analyses")
    tabs = st.tabs([
        f"{AGENT_ICONS.get(r.agent_type, '🔍')} {r.agent_type.value.replace('_', ' ').title()}"
        for r in results
    ])
    for tab, result in zip(tabs, results):
        with tab:
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"**Focus:** {result.sub_topic}")
            with col2:
                st.markdown(f"**Confidence:** {render_confidence(result.confidence)}",
                            unsafe_allow_html=True)
            st.divider()
            prose, citation_data = parse_citation_block(result.content)
            st.markdown(prose)
            render_citation_block(citation_data)


def render_fact_check(fact_check_result):
    st.subheader("🔎 Fact-Check Review")
    st.markdown(f"**Confidence:** {render_confidence(fact_check_result.confidence)}",
                unsafe_allow_html=True)
    st.divider()
    prose, citation_data = parse_citation_block(fact_check_result.content)
    st.markdown(prose)
    render_citation_block(citation_data)


def render_synthesis(synthesis):
    st.subheader("Synthesized Report")
    st.markdown(synthesis)


def render_supplementary(supplementary):
    st.divider()
    st.subheader("Supplementary Research (Separate from Synthesis)")
    st.caption("These sections are independent reference materials not included in the main synthesis.")
    tab_research, tab_gov = st.tabs([
        "🎓 Peer-Reviewed Research",
        "🏛️ Government Documents",
    ])
    with tab_research:
        r = supplementary["research"]
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"**Focus:** {r.sub_topic}")
        with col2:
            st.markdown(f"**Confidence:** {render_confidence(r.confidence)}",
                        unsafe_allow_html=True)
        st.divider()
        st.markdown(r.content)
    with tab_gov:
        g = supplementary["government"]
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"**Focus:** {g.sub_topic}")
        with col2:
            st.markdown(f"**Confidence:** {render_confidence(g.confidence)}",
                        unsafe_allow_html=True)
        st.divider()
        st.markdown(g.content)


def render_economics(economics):
    """Render economics pipeline results when active."""
    st.divider()
    st.subheader("📊 Economic Analysis (Conditional Pipeline)")
    st.caption(
        "Economics pipeline activated — detailed economic data and multi-framework "
        "policy analysis detected as relevant to this topic."
    )
    tab_data, tab_policy = st.tabs([
        "📊 Economic Data & Indicators",
        "💹 Economic Policy Analysis",
    ])
    with tab_data:
        d = economics["data"]
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"**Focus:** {d.sub_topic}")
        with col2:
            st.markdown(f"**Confidence:** {render_confidence(d.confidence)}",
                        unsafe_allow_html=True)
        st.divider()
        st.markdown(d.content)
    with tab_policy:
        p = economics["policy"]
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"**Focus:** {p.sub_topic}")
        with col2:
            st.markdown(f"**Confidence:** {render_confidence(p.confidence)}",
                        unsafe_allow_html=True)
        st.divider()
        st.markdown(p.content)


def render_meta_review_history(meta_reviews):
    """Render past meta review iterations as collapsed expanders."""
    if not meta_reviews:
        return
    for i, review in enumerate(meta_reviews):
        with st.expander(
            f"Review Iteration {i + 1} — Quality: {review.confidence:.0%}",
            expanded=False,
        ):
            st.markdown(review.critique)
            if review.adjustments:
                st.markdown("**Auto-detected Adjustments:**")
                for adj in review.adjustments:
                    st.markdown(
                        f"- **{adj['agent_type']}**: {adj['issue']}\n"
                        f"  - *Guidance:* {adj['guidance']}"
                    )


def render_sources(sources_data):
    """Render the sources tab with most-referenced and singular-critical lists."""
    st.divider()
    st.subheader("🔗 Source Analysis")
    st.caption("Sources extracted from all agent outputs, ranked by cross-agent reference frequency.")

    most_ref = sources_data.get("most_referenced", [])
    singular = sources_data.get("singular_critical", [])

    tab_top, tab_singular = st.tabs([
        f"🔗 Most Referenced ({len(most_ref)})",
        f"⚠️ Singular Critical Sources ({len(singular)})",
    ])

    with tab_top:
        if not most_ref:
            st.info("No cross-referenced sources were identified.")
        else:
            st.caption("Sources cited by multiple agents — the backbone of this analysis.")
            for i, src in enumerate(most_ref, 1):
                ref_count = src.get("reference_count", 0)
                ref_by = src.get("referenced_by", [])
                src_type = src.get("type", "other").replace("_", " ").title()
                url = src.get("url", "")
                name = src.get("name", "Unknown")

                with st.container():
                    col1, col2, col3 = st.columns([3, 1, 1])
                    with col1:
                        if url:
                            st.markdown(f"**{i}. [{name}]({url})**")
                        else:
                            st.markdown(f"**{i}. {name}**")
                    with col2:
                        st.caption(src_type)
                    with col3:
                        st.markdown(f"**{ref_count}** agents")

                    if ref_by:
                        agents_str = ", ".join(
                            a.replace("_", " ").title() for a in ref_by
                        )
                        st.caption(f"Referenced by: {agents_str}")
                    st.markdown("---")

    with tab_singular:
        if not singular:
            st.info("No singular critical sources were identified — all key claims have multiple backing sources.")
        else:
            st.caption(
                "Sources that are the **sole backing** for a critical claim. "
                "If these sources were removed or discredited, the corresponding "
                "finding would have no support."
            )
            for i, src in enumerate(singular, 1):
                name = src.get("name", "Unknown")
                url = src.get("url", "")
                claim = src.get("claim", "")
                agent = src.get("agent", "").replace("_", " ").title()
                src_type = src.get("type", "other").replace("_", " ").title()

                with st.container():
                    if url:
                        st.markdown(f"**{i}. [{name}]({url})** — _{src_type}_")
                    else:
                        st.markdown(f"**{i}. {name}** — _{src_type}_")

                    st.markdown(f"> {claim}")
                    st.caption(f"Claimed by: {agent} agent")
                    st.markdown("---")


def render_results_block(state):
    """Render all analysis results from the current pipeline state."""
    # Historical pipeline has a different display structure
    if state.get("mode") == "historical":
        render_historical_results(state)
        return

    render_sub_topics(state["sub_topics"], state.get("tangential_topics"))
    render_agent_results(state["agent_results"])
    st.divider()
    render_fact_check(state["fact_check"])
    st.divider()
    render_synthesis(state["synthesis"])

    # Sources tab (if available)
    if state.get("sources"):
        render_sources(state["sources"])

    render_supplementary(state["supplementary"])

    # Economics pipeline (conditional)
    if state.get("economics_active") and state.get("economics"):
        render_economics(state["economics"])


def render_historical_results(state):
    """Render historical pipeline results with dedicated layout."""
    context_detail = state.get("context_detail", {})
    anchor_year = state.get("anchor_year", "Unknown")

    st.markdown(f"### Historical Analysis (Anchor Year: ~{anchor_year})")

    # Ordered historical agent display
    HISTORICAL_AGENT_ORDER = [
        ("Historical Anchor", AgentType.HISTORICAL_ANCHOR, "🏛️"),
        ("Era Context", AgentType.ERA_CONTEXT, "🕰️"),
        ("Primary Sources", AgentType.PRIMARY_SOURCE, "📜"),
        ("Causal Chain", AgentType.CAUSAL_CHAIN, "🔗"),
        ("Modern Impact", AgentType.MODERN_IMPACT, "📡"),
        ("Counterfactual Analysis", AgentType.COUNTERFACTUAL, "🔀"),
        ("Ripple Timeline", AgentType.RIPPLE_TIMELINE, "🌊"),
        ("Scholarly Consensus", AgentType.SCHOLARLY_CONSENSUS, "🎓"),
        ("Perspectives", AgentType.PERSPECTIVES, "🔄"),
        ("Government Documents", AgentType.GOVERNMENT_DOCS, "🏛️"),
        ("Research Review", AgentType.RESEARCH_REVIEW, "🎓"),
    ]

    # Create tabs for each agent
    tab_labels = [f"{icon} {name}" for name, _, icon in HISTORICAL_AGENT_ORDER]
    tabs = st.tabs(tab_labels)

    results_by_type = {r.agent_type: r for r in state["agent_results"]}

    for tab, (name, agent_type, icon) in zip(tabs, HISTORICAL_AGENT_ORDER):
        with tab:
            result = results_by_type.get(agent_type)
            if result:
                conf = result.confidence
                color = "#00cc66" if conf >= 0.7 else "#ffaa00" if conf >= 0.5 else "#ff4444"
                st.markdown(
                    f"<div style='display:flex;align-items:center;gap:8px;'>"
                    f"<span style='font-size:1.5em'>{icon}</span>"
                    f"<strong>{name}</strong>"
                    f"<span style='color:{color};margin-left:auto;'>"
                    f"Confidence: {conf:.0%}</span></div>",
                    unsafe_allow_html=True)
                st.markdown(result.content)

                # Citation verification summary
                if result.citations:
                    verified = sum(1 for c in result.citations if c.get("status") == "verified")
                    unverified = sum(1 for c in result.citations if c.get("status") == "plausible_unverified")
                    fab_risk = sum(1 for c in result.citations if c.get("status") == "fabrication_risk")
                    with st.expander(f"Citations ({verified} verified, {unverified} unverified, {fab_risk} flagged)"):
                        for c in result.citations:
                            status_icon = {"verified": "✅", "plausible_unverified": "🟡",
                                           "fabrication_risk": "🔴"}.get(c.get("status", ""), "❓")
                            st.markdown(f"{status_icon} {c.get('claim_summary', 'No summary')}")
            else:
                st.info(f"No results from {name} agent")

    st.divider()

    # Fact-check
    if state.get("fact_check"):
        render_fact_check(state["fact_check"])
        st.divider()

    # Synthesis (the four-section historical brief)
    render_synthesis(state["synthesis"])

    # Sources
    if state.get("sources"):
        render_sources(state["sources"])

    # Supplementary
    if state.get("supplementary"):
        render_supplementary(state["supplementary"])

    # Economics (conditional)
    if state.get("economics_active") and state.get("economics"):
        render_economics(state["economics"])


def build_download_report(state):
    """Build the full downloadable markdown report."""
    topic = state.get("raw_topic", state["topic"])
    full_report = f"# DeepDive Analysis: {topic}\n\n"

    for r in state["agent_results"]:
        full_report += f"## {r.agent_type.value.replace('_', ' ').title()} Analysis\n"
        full_report += f"Focus: {r.sub_topic}\n"
        full_report += f"Confidence: {r.confidence:.0%}\n\n"
        full_report += f"{r.content}\n\n---\n\n"

    full_report += f"## Fact-Check Review\n"
    full_report += f"Confidence: {state['fact_check'].confidence:.0%}\n\n"
    full_report += f"{state['fact_check'].content}\n\n---\n\n"

    full_report += f"## Synthesis\n\n{state['synthesis']}\n\n---\n\n"

    if state["meta_reviews"]:
        full_report += f"## Meta Reviews\n\n"
        for i, review in enumerate(state["meta_reviews"]):
            full_report += f"### Iteration {i + 1}\n"
            full_report += f"Quality: {review.confidence:.0%}\n\n"
            full_report += f"{review.critique}\n\n"
            if review.adjustments:
                full_report += "**Adjustments:**\n"
                for adj in review.adjustments:
                    full_report += f"- {adj['agent_type']}: {adj['issue']}\n"
            full_report += "\n---\n\n"

    # Sources
    sources = state.get("sources")
    if sources:
        full_report += f"## Source Analysis\n\n"
        full_report += f"### Most Referenced Sources\n\n"
        for src in sources.get("most_referenced", []):
            url_part = f" — [{src.get('url')}]({src.get('url')})" if src.get("url") else ""
            refs = ", ".join(a.replace("_", " ").title() for a in src.get("referenced_by", []))
            full_report += (
                f"- **{src.get('name', 'Unknown')}** ({src.get('type', 'other').replace('_', ' ').title()})"
                f"{url_part} — Referenced by {src.get('reference_count', 0)} agents: {refs}\n"
            )
        full_report += f"\n### Singular Critical Sources\n\n"
        for src in sources.get("singular_critical", []):
            url_part = f" — [{src.get('url')}]({src.get('url')})" if src.get("url") else ""
            full_report += (
                f"- **{src.get('name', 'Unknown')}** ({src.get('type', 'other').replace('_', ' ').title()})"
                f"{url_part}\n  > {src.get('claim', '')}\n  Claimed by: {src.get('agent', '').replace('_', ' ').title()}\n\n"
            )
        full_report += "---\n\n"

    supplementary = state["supplementary"]
    full_report += f"## Supplementary: Peer-Reviewed Research\n"
    full_report += f"Confidence: {supplementary['research'].confidence:.0%}\n\n"
    full_report += f"{supplementary['research'].content}\n\n---\n\n"

    full_report += f"## Supplementary: Government Documents\n"
    full_report += f"Confidence: {supplementary['government'].confidence:.0%}\n\n"
    full_report += f"{supplementary['government'].content}\n\n---\n\n"

    # Economics pipeline (conditional)
    economics = state.get("economics", {})
    if state.get("economics_active") and economics:
        full_report += f"## Economic Analysis (Conditional Pipeline)\n\n"
        if economics.get("data"):
            full_report += f"### Economic Data & Indicators\n"
            full_report += f"Confidence: {economics['data'].confidence:.0%}\n\n"
            full_report += f"{economics['data'].content}\n\n---\n\n"
        if economics.get("policy"):
            full_report += f"### Economic Policy Analysis\n"
            full_report += f"Confidence: {economics['policy'].confidence:.0%}\n\n"
            full_report += f"{economics['policy'].content}\n\n"

    return full_report


def _add_markdown_content_to_docx(doc, content):
    """Convert markdown content to docx paragraphs, stripping JSON code blocks."""
    # Pre-strip all fenced code blocks (including malformed ones) before line processing
    content = re.sub(r'```(?:json)?\s*\{.*?\}\s*```', '', content, flags=re.DOTALL)
    in_code_block = False
    for line in content.split("\n"):
        stripped = line.strip()
        # Toggle code fence state
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        # Skip lines inside code blocks (JSON citations, etc.)
        if in_code_block:
            continue
        # Skip stray JSON lines that leaked through
        if stripped.startswith('{') and stripped.endswith('}'):
            continue
        if stripped.startswith('"') and stripped.endswith(','):
            continue
        if not stripped:
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("---"):
            doc.add_paragraph("—" * 40)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        else:
            # Strip markdown bold/italic markers
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            doc.add_paragraph(clean)


def build_docx_report(state):
    """Build a .docx version of the report and return as bytes."""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    topic = state["topic"]
    doc.add_heading(f'DeepDive Analysis: {topic}', level=0)

    for r in state["agent_results"]:
        doc.add_heading(f'{r.agent_type.value.replace("_", " ").title()} Analysis', level=1)
        doc.add_paragraph(f'Focus: {r.sub_topic}')
        doc.add_paragraph(f'Confidence: {r.confidence:.0%}')
        _add_markdown_content_to_docx(doc, r.content)

    doc.add_heading('Fact-Check Review', level=1)
    doc.add_paragraph(f'Confidence: {state["fact_check"].confidence:.0%}')
    _add_markdown_content_to_docx(doc, state["fact_check"].content)

    doc.add_heading('Synthesis', level=1)
    _add_markdown_content_to_docx(doc, state["synthesis"])

    if state["meta_reviews"]:
        doc.add_heading('Meta Reviews', level=1)
        for i, review in enumerate(state["meta_reviews"]):
            doc.add_heading(f'Iteration {i + 1} — Quality: {review.confidence:.0%}', level=2)
            _add_markdown_content_to_docx(doc, review.critique)

    supplementary = state["supplementary"]
    doc.add_heading('Supplementary: Peer-Reviewed Research', level=1)
    _add_markdown_content_to_docx(doc, supplementary['research'].content)

    doc.add_heading('Supplementary: Government Documents', level=1)
    _add_markdown_content_to_docx(doc, supplementary['government'].content)

    economics = state.get("economics", {})
    if state.get("economics_active") and economics:
        doc.add_heading('Economic Analysis', level=1)
        if economics.get("data"):
            doc.add_heading('Economic Data & Indicators', level=2)
            _add_markdown_content_to_docx(doc, economics['data'].content)
        if economics.get("policy"):
            doc.add_heading('Economic Policy Analysis', level=2)
            _add_markdown_content_to_docx(doc, economics['policy'].content)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def save_parse_result_log(raw_topic, parse_result, engineered_topic):
    """Log ParseResult to disk for QC review. Not displayed in the dashboard."""
    os.makedirs(RESULTS_DIR, exist_ok=True)
    slug = re.sub(r'[^a-z0-9]+', '_', raw_topic.lower().strip())[:60].strip('_')
    log_entry = {
        "timestamp": datetime.now().isoformat(),
        "raw_topic": raw_topic,
        "parse_result": asdict(parse_result),
        "engineered_topic": engineered_topic,
    }
    log_path = os.path.join(RESULTS_DIR, f"{slug}_parser_qc.json")
    with open(log_path, "w", encoding="utf-8") as f:
        json.dump(log_entry, f, indent=2, default=str)


def auto_save_results(state):
    """Save the latest results to disk so they can be read outside Streamlit."""
    os.makedirs(RESULTS_DIR, exist_ok=True)
    topic = state.get("raw_topic", state.get("topic", "unknown"))
    slug = re.sub(r'[^a-z0-9]+', '_', topic.lower().strip())[:60].strip('_')

    # Save the full markdown report
    report = build_download_report(state)
    report_path = os.path.join(RESULTS_DIR, f"{slug}.md")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report)

    # Also save as "latest.md" for easy access
    latest_path = os.path.join(RESULTS_DIR, "latest.md")
    with open(latest_path, "w", encoding="utf-8") as f:
        f.write(report)


# ── Dashboard tabs: Logs and Vector DB ──────────────────────────────

def render_logs_tab():
    """Display saved analysis log files."""
    st.subheader("📄 Analysis Logs")
    st.caption("Previously saved analysis reports.")

    log_files = sorted(
        globlib.glob(os.path.join(RESULTS_DIR, "*.md")),
        key=os.path.getmtime,
        reverse=True,
    )

    if not log_files:
        st.info("No analysis logs found. Run an analysis first.")
        return

    for i, filepath in enumerate(log_files):
        filename = os.path.basename(filepath)
        if filename == "latest.md":
            continue
        mod_time = os.path.getmtime(filepath)
        from datetime import datetime
        mod_str = datetime.fromtimestamp(mod_time).strftime("%Y-%m-%d %H:%M")
        display_name = filename.replace("_", " ").replace(".md", "").title()

        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()

        total_len = len(content)
        preview_len = max(200, int(total_len * 0.05))
        expanded_len = max(400, int(total_len * 0.08))

        # Strip leading H1 for body content
        preview_content = content
        if preview_content.startswith("# "):
            preview_content = preview_content.split("\n", 1)[-1]

        # Step 1: Title + perpetual actions row
        col_title, col_open, col_gdrive = st.columns([4, 1, 1])
        with col_title:
            st.markdown(f"**{display_name}** — {mod_str}")
        with col_open:
            # Download the full report as a styled HTML file
            import re as _re
            escaped = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            html_lines = []
            for line in escaped.split("\n"):
                if line.startswith("### "):
                    html_lines.append(f"<h3>{line[4:]}</h3>")
                elif line.startswith("## "):
                    html_lines.append(f"<h2>{line[3:]}</h2>")
                elif line.startswith("# "):
                    html_lines.append(f"<h1>{line[2:]}</h1>")
                elif line.startswith("---"):
                    html_lines.append("<hr>")
                elif line.startswith("- "):
                    html_lines.append(f"<li>{line[2:]}</li>")
                elif line.strip() == "":
                    html_lines.append("<br>")
                else:
                    line = _re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', line)
                    html_lines.append(f"<p>{line}</p>")
            html_full = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{display_name}</title>
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
max-width: 900px; margin: 2rem auto; padding: 0 1.5rem; line-height: 1.7;
color: #e0e0e0; background: #0a0a0a; }}
h1,h2,h3 {{ color: #fff; border-bottom: 1px solid #333; padding-bottom: .3em; }}
h1 {{ font-size: 1.8rem; }} h2 {{ font-size: 1.4rem; }} h3 {{ font-size: 1.1rem; }}
hr {{ border: none; border-top: 1px solid #333; margin: 2rem 0; }}
blockquote {{ border-left: 3px solid #39ff14; padding-left: 1rem; color: #aaa; }}
code {{ background: #1a1a1a; padding: .15em .4em; border-radius: 3px; font-size: .9em; }}
strong {{ color: #fff; }}
li {{ margin-left: 1.5rem; margin-bottom: .3rem; }}
</style></head><body>
""" + "\n".join(html_lines) + "\n</body></html>"
            # Use JS window.open + document.write to pop out in a new tab
            # Escape for safe embedding in JS string
            import json as _json
            js_escaped = _json.dumps(html_full)
            component_html = f"""
            <script>
            function openFullReport_{i}() {{
                var w = window.open('', '_blank');
                if (w) {{
                    w.document.open();
                    w.document.write({js_escaped});
                    w.document.close();
                }} else {{
                    alert('Pop-up blocked. Please allow pop-ups for this site.');
                }}
            }}
            </script>
            <button onclick="openFullReport_{i}()"
                style="background:#333;color:#fff;border:1px solid #555;
                padding:.35rem .8rem;border-radius:4px;cursor:pointer;
                font-size:.8rem;width:100%;">Open Full</button>
            """
            st.components.v1.html(component_html, height=40)
        with col_gdrive:
            if is_drive_configured():
                if st.button("Create DOCX", key=f"gdrive_{filename}"):
                    try:
                        # Build DOCX from markdown content
                        docx_doc = Document()
                        style = docx_doc.styles['Normal']
                        style.font.name = 'Calibri'
                        style.font.size = Pt(11)

                        _add_markdown_content_to_docx(docx_doc, content)

                        buf = io.BytesIO()
                        docx_doc.save(buf)
                        buf.seek(0)

                        # Build filename: LOG_YYYY-MM-DD_truncated_title.docx
                        from datetime import datetime as _dt
                        log_date = _dt.fromtimestamp(mod_time).strftime("%Y-%m-%d")
                        truncated_title = display_name[:40].strip()
                        drive_filename = f"LOG_{log_date}_{truncated_title}.docx"
                        drive_filename = re.sub(r'[^a-zA-Z0-9_\-. ]', '', drive_filename)

                        result = upload_docx_to_drive(buf.getvalue(), drive_filename)
                        link = result.get("webViewLink", "")
                        st.success(f"Created: [{result.get('name')}]({link})")
                    except Exception as e:
                        st.error(f"Failed: {e}")
            else:
                st.caption("Drive not configured")

        # Step 2: ~5% preview (click to expand)
        with st.expander("Preview", expanded=False):
            st.markdown(preview_content[:preview_len])
            if total_len > preview_len:
                st.caption(f"~5% ({preview_len} of {total_len} chars)")

            # Step 3: ~8% expanded (nested)
            if total_len > preview_len:
                with st.expander("Show more", expanded=False):
                    st.markdown(preview_content[:expanded_len])
                    if total_len > expanded_len:
                        st.caption(f"~8% ({expanded_len} of {total_len} chars)")

        st.divider()


def render_vector_db_tab():
    """Display plain-text corrections stored in the vector database."""
    st.subheader("🧠 Vector Database — Stored Corrections")
    st.caption("Correction patterns stored from user feedback. These influence future analyses.")

    store = CorrectionStore()
    corrections = store.get_all_corrections(limit=200)

    if not corrections:
        st.info("No corrections stored yet. Corrections are saved when you provide feedback during meta review.")
        return

    st.markdown(f"**{len(corrections)} correction(s) stored**")

    for i, c in enumerate(corrections, 1):
        agent = c.get("agent_type", "unknown")
        topic = c.get("topic", "N/A")
        issue = c.get("issue", "N/A")
        user_fb = c.get("user_feedback", "N/A")
        guidance = c.get("optimized_guidance", "N/A")
        effective = c.get("effective", "unknown")
        timestamp = c.get("timestamp", "")

        eff_icon = {"yes": "✅", "no": "❌", "unknown": "❓"}.get(effective, "❓")

        with st.expander(
            f"{i}. [{agent}] {issue[:80]}{'...' if len(issue) > 80 else ''} {eff_icon}",
            expanded=False,
        ):
            st.markdown(f"**Agent:** `{agent}`")
            st.markdown(f"**Topic:** {topic}")
            st.markdown(f"**Issue:** {issue}")
            st.markdown(f"**User Feedback:** {user_fb}")
            st.markdown(f"**Optimized Guidance:** {guidance}")
            st.markdown(f"**Effective:** {eff_icon} {effective}")
            if timestamp:
                st.caption(f"Stored: {timestamp}")

    st.divider()
    if st.button("Clear All Corrections", type="secondary"):
        store.clear_all()
        st.success("Vector database cleared.")
        st.rerun()


# ── Main application flow ──────────────────────────────────────────

def main():
    init_session_state()
    render_header()

    # ── Sidebar tabs for logs and vector DB ──
    main_tab, logs_tab, vectordb_tab = st.tabs([
        "🔍 Analysis",
        "📄 Logs",
        "🧠 Vector DB",
    ])

    with logs_tab:
        render_logs_tab()

    with vectordb_tab:
        render_vector_db_tab()

    with main_tab:
        # ── INPUT PHASE ──
        if st.session_state.phase == "input":
            topic, analyze = render_input()

            if analyze and topic:
                st.session_state.phase = "analyzing"
                st.session_state.orchestrator = Orchestrator()
                st.session_state._topic = topic
                st.rerun()
            elif analyze and not topic:
                st.warning("Please enter a topic to analyze.")

        # ── ANALYZING PHASE ──
        elif st.session_state.phase == "analyzing":
            topic = st.session_state._topic
            if not topic:
                st.session_state.phase = "input"
                st.rerun()
            orchestrator = st.session_state.orchestrator

            with st.status("Running analysis pipeline...", expanded=True) as status:
                try:
                    # Step 0: Engineer the prompt
                    st.write("Engineering research prompt from your topic...")
                    engineered_topic = orchestrator.engineer_prompt(topic)
                    # Log ParseResult for QC (not shown in dashboard)
                    if hasattr(orchestrator, 'last_parse_result'):
                        save_parse_result_log(topic, orchestrator.last_parse_result, engineered_topic)
                    st.write("Research prompt optimized")
                    with st.expander("View engineered prompt", expanded=False):
                        st.markdown(f"**Your input:** {topic}")
                        st.markdown(f"**Engineered prompt:** {engineered_topic}")

                    # Run the full pipeline through the single correct entry point
                    # This activates: anchor detection, pre-verification news fetch,
                    # URL prefetching, seed evidence injection, parallel dispatch,
                    # economics detection, fact-check, source extraction, synthesis
                    def progress_callback(agent_type, result):
                        st.write(f"✓ {agent_type.value.replace('_', ' ').title()} complete")

                    state = orchestrator.run_analysis(engineered_topic, progress_callback)
                    state["raw_topic"] = topic

                    # Generate tangential topics (separate from main pipeline)
                    st.write("Generating tangential topics...")
                    state["tangential_topics"] = orchestrator.generate_tangential_topics(
                        engineered_topic, state["sub_topics"]
                    )
                    st.write("Tangential topics generated")

                    # Ensure sources key exists for dashboard rendering
                    if "sources" not in state:
                        state["sources"] = None

                    # Step 9: Extract sources
                    st.write("Extracting and ranking sources across all agents...")
                    state["sources"] = orchestrator.extract_sources(state)
                    st.write("Source extraction complete")

                    # Step 10: Meta review
                    st.write("Running meta review...")
                    state = orchestrator.run_meta_review(state, iteration=0)
                    st.write("Meta review complete")

                    st.session_state.pipeline_state = state
                    st.session_state.phase = "meta_review"
                    st.session_state.correction_round = 0
                    auto_save_results(state)
                    status.update(label="Analysis complete — review below",
                                  state="complete", expanded=False)
                except Exception as e:
                    status.update(label="Analysis failed", state="error")
                    st.error(f"Error during analysis: {e}")
                    import traceback
                    st.code(traceback.format_exc())
                    if st.button("Start Over"):
                        st.session_state.phase = "input"
                        st.session_state.pipeline_state = None
                        st.rerun()
                    return

            st.rerun()

        # ── META REVIEW PHASE (interactive) ──
        elif st.session_state.phase == "meta_review":
            state = st.session_state.pipeline_state
            topic = state["topic"]

            st.info(f"**Analyzing:** {topic}")

            # Show all results
            render_results_block(state)

            # Show meta review
            st.divider()
            st.subheader("🧠 Meta Review — Your Input Needed")
            st.caption(
                "The Meta Agent has reviewed all outputs. Read its critique below, "
                "then provide your own corrections or approve the results."
            )

            # Show past review iterations
            if len(state["meta_reviews"]) > 1:
                render_meta_review_history(state["meta_reviews"][:-1])

            # Show latest review prominently
            latest_review = state["meta_reviews"][-1]
            round_label = f"Round {st.session_state.correction_round + 1}"

            st.markdown(f"### {round_label} Critique")
            quality_color = "🟢" if latest_review.confidence >= 0.7 else "🟡" if latest_review.confidence >= 0.5 else "🔴"
            st.markdown(f"**Overall Quality:** {quality_color} {latest_review.confidence:.0%}")

            # Accuracy verdict
            VERDICT_LABELS = {
                "accurate": ("✅", "Accurate"),
                "mostly_accurate": ("🟡", "Mostly Accurate"),
                "has_errors": ("🟠", "Has Errors"),
                "unreliable": ("🔴", "Unreliable"),
            }
            verdict = getattr(latest_review, "accuracy_verdict", "mostly_accurate")
            v_icon, v_label = VERDICT_LABELS.get(verdict, ("🟡", "Mostly Accurate"))
            st.markdown(f"**Accuracy Verdict:** {v_icon} {v_label}")
            st.markdown(latest_review.critique)

            # Downloads available during review
            st.divider()
            dl_col1, dl_col2 = st.columns(2)
            full_report = build_download_report(state)
            with dl_col1:
                st.download_button(
                    label="Download Markdown Report",
                    data=full_report,
                    file_name=f"deepdive_{topic[:30].replace(' ', '_')}.md",
                    mime="text/markdown",
                    key="dl_md_meta",
                )
            with dl_col2:
                docx_bytes = build_docx_report(state)
                st.download_button(
                    label="Download DOCX Report",
                    data=docx_bytes,
                    file_name=f"deepdive_{topic[:30].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_docx_meta",
                )
            # Google Drive upload
            if is_drive_configured():
                if st.button("Upload to Google Drive", key="gdrive_meta"):
                    try:
                        docx_data = build_docx_report(state)
                        fname = f"deepdive_{topic[:30].replace(' ', '_')}.docx"
                        result = upload_docx_to_drive(docx_data, fname)
                        link = result.get("webViewLink", "")
                        st.success(f"Uploaded to Google Drive: [{result.get('name')}]({link})")
                    except Exception as e:
                        st.error(f"Drive upload failed: {e}")
            st.divider()

            if latest_review.adjustments:
                st.markdown("**Meta Agent's Suggested Adjustments:**")
                for adj in latest_review.adjustments:
                    st.markdown(
                        f"- **{adj['agent_type']}**: {adj['issue']}\n"
                        f"  - *Suggested fix:* {adj['guidance']}"
                    )

            # User correction input
            st.divider()
            st.markdown("### Your Corrections")
            st.caption(
                "Tell us what to fix. Be as specific or general as you like — "
                "the Meta Agent will optimize your feedback into targeted prompt adjustments "
                "and store the patterns for future analyses."
            )

            user_feedback = st.text_area(
                "Your corrections (leave empty and click Approve to accept as-is)",
                height=150,
                placeholder="e.g., The perspectives agent missed the labor union angle. "
                            "The timeline is missing events from 2023. "
                            "The facts agent should focus more on economic data.",
                key=f"user_corrections_{st.session_state.correction_round}",
            )

            col_approve, col_correct, col_restart = st.columns([1, 1, 1])

            with col_approve:
                approve = st.button("✅ Approve Output", type="primary", use_container_width=True)

            with col_correct:
                correct = st.button("🔄 Apply Corrections", use_container_width=True,
                                    disabled=not user_feedback.strip())

            with col_restart:
                restart = st.button("🔁 Start Over", use_container_width=True)

            if approve:
                orchestrator = st.session_state.orchestrator
                orchestrator.approve_output(state)
                auto_save_results(state)
                st.session_state.phase = "done"
                st.rerun()

            elif correct and user_feedback.strip():
                st.session_state.phase = "correcting"
                st.session_state._user_feedback = user_feedback.strip()
                st.rerun()

            elif restart:
                st.session_state.phase = "input"
                st.session_state.pipeline_state = None
                st.rerun()

        # ── CORRECTING PHASE ──
        elif st.session_state.phase == "correcting":
            state = st.session_state.pipeline_state
            orchestrator = st.session_state.orchestrator
            user_feedback = st.session_state._user_feedback

            with st.status("Applying corrections...", expanded=True) as status:
                try:
                    st.write("🧠 Optimizing your feedback into targeted prompt adjustments...")
                    state = orchestrator.apply_user_corrections(state, user_feedback)

                    st.write("🧠 Re-running meta review on corrected output...")
                    iteration = len(state["meta_reviews"])
                    state = orchestrator.run_meta_review(state, iteration=iteration)

                    st.session_state.pipeline_state = state
                    st.session_state.correction_round += 1
                    st.session_state.phase = "meta_review"
                    auto_save_results(state)
                    status.update(label="Corrections applied — review updated results",
                                  state="complete", expanded=False)
                except Exception as e:
                    status.update(label="Correction failed", state="error")
                    st.error(f"Error applying corrections: {e}")
                    st.session_state.phase = "meta_review"

            st.rerun()

        # ── DONE PHASE ──
        elif st.session_state.phase == "done":
            state = st.session_state.pipeline_state
            topic = state["topic"]

            st.success(f"Analysis complete: **{topic}**")

            if st.session_state.correction_round > 0:
                st.caption(
                    f"{st.session_state.correction_round} correction round(s) applied. "
                    f"Patterns stored in vector database for future analyses."
                )

            # Show final results
            render_results_block(state)

            # Show meta review history
            if state["meta_reviews"]:
                st.divider()
                st.subheader("🧠 Meta Review History")
                render_meta_review_history(state["meta_reviews"])

            # Downloads
            st.divider()
            st.subheader("Downloads")
            col_md, col_docx = st.columns(2)

            full_report = build_download_report(state)
            with col_md:
                st.download_button(
                    label="Download Markdown Report",
                    data=full_report,
                    file_name=f"deepdive_{topic[:30].replace(' ', '_')}.md",
                    mime="text/markdown",
                )

            with col_docx:
                docx_bytes = build_docx_report(state)
                st.download_button(
                    label="Download DOCX Report",
                    data=docx_bytes,
                    file_name=f"deepdive_{topic[:30].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            # Google Drive upload
            if is_drive_configured():
                if st.button("Upload to Google Drive", key="gdrive_done"):
                    try:
                        docx_data = build_docx_report(state)
                        fname = f"deepdive_{topic[:30].replace(' ', '_')}.docx"
                        result = upload_docx_to_drive(docx_data, fname)
                        link = result.get("webViewLink", "")
                        st.success(f"Uploaded to Google Drive: [{result.get('name')}]({link})")
                    except Exception as e:
                        st.error(f"Drive upload failed: {e}")

            # New analysis button
            st.divider()
            if st.button("Start New Analysis"):
                st.session_state.phase = "input"
                st.session_state.pipeline_state = None
                st.session_state.correction_round = 0
                st.rerun()


if __name__ == "__main__":
    main()
